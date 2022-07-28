import pyppeteer
import json
import re, os,time
import xlwt
import logging
import docx
from docx.opc.constants import CONTENT_TYPE as CT
from docx.package import Package
import zipfile
import xlsxwriter

location = []
not_isdir = []
not_files = []
all_title_name = ['【国药准字】',
                # '【药品名称】',
                '【通用名称】',
                '【商品名称】',
                '【英文名称】',
                '【成份】',
                '【主要成份】',
                '【性状】',
                '【适应症】',
                '【规格】',
                '【用法用量】',
                '【不良反应】',
                '【禁忌】',
                '【注意事项】',
                '【孕妇及哺乳期妇女用药】',
                '【儿童用药】',
                '【老年用药】',
                '【药物相互作用】',
                '【药物过量】',
                '【药理毒理】',
                '【药代动力学】',
                '【贮藏】',
                '【包装】',
                '【有效期】',
                '【生产企业】',
                '[国药准字]',
                  '[通用名称]',
                  # '[药品名称]',
                  '[商品名称]',
                  '[英文名称]',
                  '[成份]',
                  '[主要成份]',
                  '[性状]',
                  '[适应症]',
                  '[规格]',
                  '[用法用量]',
                  '[不良反应]',
                  '[禁忌]',
                  '[注意事项]',
                  '[孕妇及哺乳期妇女用药]',
                  '[儿童用药]',
                  '[老年用药]',
                  '[药物相互作用]',
                  '[药物过量]',
                  '[药理毒理]',
                  '[药代动力学]',
                  '[贮藏]',
                  '[包装]',
                  '[有效期]',
                  '[生产企业]'
                  ]

async def get_html(url):
    browser = await pyppeteer.launch(headless=True, args=['--no-sandbox'])
    page = await  browser.newPage()
    res = await page.goto(url, options={'timeout': 3000})
    data = await page.content()
    title = await page.title()
    resp_cookies = await page.cookies()  # cookie
    resp_headers = res.headers  # 响应头
    resp_status = res.status  # 响应状态
    # print(data)
    # print(title)
    # print(resp_headers)
    # print(resp_status)
    return data

def iterate_files():
    file_path = os.path.join(os.getcwd(), '无法识别/')
    # 遍历文件夹
    files = []
    for root, dirs, files in os.walk(file_path):
        # print(root)  # 当前目录路径
        # print(dirs)  # 当前路径下所有子目录
        # print(files)  # 当前路径下所有非目录子文件
        print(files)
        files = files
        i = 0
    for file_name in files:
        # i=i+1
        # if i == 10:
        #     break
        if file_name != '.DS_Store':
            data = decode_docx(file_name)
            if len(data) > 0:
                location.append(data)
    return location
#解析docx说明书
def decode_docx(file_name):
    file_path = os.path.join(os.getcwd(), '无法识别/')
    file_dir = file_path[:-1]
    if not os.path.exists(file_dir):
        logging.info("Mkdir '无法识别/'.")
        os.mkdir(file_dir)

    file_path = file_dir + '/' + file_name
    print('file_path:'+file_path)
    # 获得文件大小
    sz = os.path.getsize(file_path)
    if sz < 2500:
        return []
    if os.path.isdir(file_path):
        not_isdir.append(file_path)
        return []
    elif zipfile.is_zipfile(file_path):
        print()
        # 获取content_type
        try:
            type = Package.open(file_path)
            type.main_document_part
        except KeyError:
            print('Error: 读取文件document_part失败')
            not_isdir.append(file_path)
            return []
        except FileNotFoundError:
            print('Error: No such file or directory失败')
            not_isdir.append(file_path)
            return []
        document_part = Package.open(file_path).main_document_part
        print('获取content_type:' + document_part.content_type)
        # 获取文章全部内容
        # if not os.path.exists(file_path):
        #     not_files.append(file_path)
        if document_part.content_type != CT.WML_DOCUMENT_MAIN:
            not_files.append(file_path)
            return []
        else:
            doc = docx.Document(file_path)
        # 一级标题
        heading_1 = []
        for p in doc.paragraphs:
            if p.style.name == 'Heading 1':
                # print('一级标题' + p.text)
                heading_1.append(p.text)
        # 二级标题
        heading_2 = []
        for p in doc.paragraphs:
            if p.style.name == 'Heading 2':
                # print('二级标题' + p.text)
                heading_2.append(p.text)
        # 所有内容
        normal_type = []
        normal_content = []
        normal_content.append(file_name)
        doc_content = doc.paragraphs
        # 不良反应
        i_1 = 0
        i_2 = 0
        # 解析 具体title
        # decode_context(p,doc_content,normal_type,normal_content)

        # 解析 全量Title
        decode_content_all(p.style.name,doc_content,normal_content)

        # 使用json.dumps()方法转为json格式数据
        # 注意：默认会转为二进制数据，使用 ensure_ascii=False 设置不转为二进制
        json_data = json.dumps(location, ensure_ascii=False)
        print('dumps_content:' + json_data)
        return normal_content
        # return location
    else:
        print(file_path)
        not_isdir.append(file_path)
        return []

list_error = []

def decode_content_all(style_name,doc_content,normal_content):
    # 不良反应
    i_1 = 0
    i_2 = 0
    content = ''
    if style_name == 'Normal':
        for i in range(0, len(doc_content)):
            if '】' in doc_content[i].text:
                style_name_arr = doc_content[i].text.split('】')[0] + '】'
            else:
                style_name_arr = doc_content[i].text
            #剔除空格
            style_name_arr = style_name_arr.replace(' ', '')
            if doc_content[i].text.replace(' ', '') == '【药品名称】' or doc_content[i].text.replace(' ', '') == '[药品名称]':
                content_1 = doc_content[i + 1].text
                if '商品名称' in content_1 or '英文名称' in content_1:
                    text = content_1.split('\n')
                    content_array_1 = text[0].split('：')
                    content_array_2 = text[1].split('：')
                    content_array_3 = text[1].split('：')
                    if content_array_1[0].replace(' ', '') == '通用名称' or content_array_1[0].replace(' ', '') == '通用名':
                        normal_content.append(content_array_1[1])
                    else:
                        normal_content.append('')
                    if content_array_2[0].replace(' ', '') == '商品名称' or content_array_2[0].replace(' ', '') == '商品名':
                        normal_content.append(content_array_2[1])
                    else:
                        normal_content.append('')
                    if content_array_2[0].replace(' ', '') == '英文名称' or content_array_2[0].replace(' ', '') == '英文名':
                        normal_content.append(content_array_2[1])
                    else:
                        if len(content_array_3) >= 1:
                            normal_content.append(content_array_3[1])
                        else:
                            normal_content.append('')
                else:
                    content_2 = doc_content[i + 2].text
                    content_3 = doc_content[i + 3].text
                    content_array_1 = content_1.split('：')
                    content_array_2 = content_2.split('：')
                    content_array_3 = content_3.split('：')
                    if content_array_1[0].replace(' ', '') == '通用名称' or content_array_1[0].replace(' ', '') == '通用名':
                        normal_content.append(content_array_1[1])
                    else:
                        normal_content.append('')
                    if content_array_2[0].replace(' ', '') == '商品名称' or content_array_2[0].replace(' ', '') == '商品名':
                        normal_content.append(content_array_2[1])
                    else:
                        normal_content.append('')
                    if content_array_2[0].replace(' ', '') == '英文名称' or content_array_2[0].replace(' ', '') == '英文名':
                        normal_content.append(content_array_2[1])
                    else:
                        if len(content_array_3) > 1:
                            normal_content.append(content_array_3[1])
                        elif len(content_array_3) == 1:
                            normal_content.append(content_array_3[0])
                        else:
                            normal_content.append('')

            if style_name_arr in all_title_name:
                if i_1 == 0:
                    title_name = doc_content[i].text
                    print('title_name:' + title_name)
                    # 获取 适应症开始节点
                    i_1 = i + 1
                elif title_name != doc_content[i].text:
                    # 获取 适应症结束节点
                    i_2 = i
                    # 获取 适应症
                    for j in range(i_1, i_2):
                        content = content + doc_content[j].text
                    normal_content.append(title_name+content)
                    # 获取 不良反应 开始节点
                    i_1 = i + 1
                    content = ''
                    title_name = doc_content[i].text




def decode_context(p,doc_content,normal_type,normal_content):
    for i in range(0, len(doc_content)):
        if p.style.name == 'Normal':
            if doc_content[i].text == '【药品名称】':
                normal_type.append(doc_content[i].text)
                content_1 = doc_content[i + 1].text
                if '商品名称' in content_1 or '英文名称' in content_1:
                    text = content_1.split('\n')
                    content_array_1 = text[0].split('：')
                    content_array_2 = text[1].split('：')
                    content_array_3 = text[1].split('：')
                    if content_array_1[0] == '通用名称':
                        normal_content.append(content_array_1[1])
                    else:
                        normal_content.append('')
                    if content_array_2[0] == '商品名称':
                        normal_content.append(content_array_2[1])
                    else:
                        normal_content.append('')
                    if content_array_2[0] == '英文名称' or content_array_2[0] == '英文名':
                        normal_content.append(content_array_2[1])
                    else:
                        if len(content_array_3) >= 1:
                            normal_content.append(content_array_3[1])
                        else:
                            normal_content.append('')
                else:
                    content_2 = doc_content[i + 2].text
                    content_3 = doc_content[i + 3].text
                    content_array_1 = content_1.split('：')
                    content_array_2 = content_2.split('：')
                    content_array_3 = content_3.split('：')
                    if content_array_1[0] == '通用名称':
                        normal_content.append(content_array_1[1])
                    else:
                        normal_content.append('')
                    if content_array_2[0] == '商品名称':
                        normal_content.append(content_array_2[1])
                    else:
                        normal_content.append('')
                    if content_array_2[0] == '英文名称' or content_array_2[0] == '英文名':
                        normal_content.append(content_array_2[1])
                    else:
                        if len(content_array_3) > 1:
                            normal_content.append(content_array_3[1])
                        elif len(content_array_3) == 1:
                            normal_content.append(content_array_3[0])
                        else:
                            normal_content.append('')
            if doc_content[i].text == '【适应症】':
                if len(doc_content[i].text) > len('【适应症】'):
                    normal_content.append(doc_content[i + 1].text)
                else:
                    # 获取 适应症开始节点
                    i_1 = i + 1
                    normal_type.append(doc_content[i].text)
                    normal_content.append(doc_content[i + 1].text)

            if doc_content[i].text == '【不良反应】':
                if len(doc_content[i].text) > len('【不良反应】'):
                    normal_content.append(doc_content[i + 1].text)
                else:
                    # 获取 适应症结束节点
                    i_2 = i - 1
                    # 获取 适应症
                    content = ''
                    for j in range(i_1, i_2):
                        content = content + doc_content[j].text
                    normal_type.append(doc_content[i].text)
                    normal_content.append(content)
                    # 获取 不良反应 开始节点
                    i_1 = i + 1
            if doc_content[i].text == '【注意事项】':
                if len(doc_content[i].text) > len('【注意事项】'):
                    normal_content.append(doc_content[i + 1].text)
                else:
                    # 获取 不良反应 结束节点
                    i_2 = i - 1
                    # 获取 不良反应
                    content = ''
                    for j in range(i_1, i_2):
                        content = content + doc_content[j].text
                    normal_type.append(doc_content[i].text)
                    normal_content.append(content)
                    # 获取 注意事项 开始节点
                    i_1 = i + 1
            if doc_content[i].text == '【禁忌】':
                if len(doc_content[i].text) > len('【禁忌】'):
                    normal_content.append(doc_content[i + 1].text)
                else:
                    # 获取 注意事项 结束节点
                    i_2 = i - 1
                    # 获取 注意事项
                    content = ''
                    for j in range(i_1, i_2):
                        content = content + doc_content[j].text
                    normal_type.append(doc_content[i].text)
                    normal_content.append(content)
                    # 获取 禁忌 开始节点
                    i_1 = i + 1
            if doc_content[i].text == '【药物过量】':
                if len(doc_content[i].text) > len('【药物过量】'):
                    normal_content.append(doc_content[i + 1].text)
                else:
                    # 获取 禁忌 结束节点
                    i_2 = i - 1
                    # 获取 禁忌
                    content = ''
                    for j in range(i_1, i_2):
                        content = content + doc_content[j].text
                    normal_type.append(doc_content[i].text)
                    normal_content.append(content)
                    # 获取 药物过量 开始节点
                    i_1 = i + 1
            if doc_content[i].text == '【药理毒理】':
                if len(doc_content[i].text) > len('【药物过量】'):
                    normal_content.append(doc_content[i + 1].text)
                else:
                    # 获取 药物过量 结束节点
                    i_2 = i - 1
                    # 获取 药物过量
                    content = ''
                    for j in range(i_1, i_2):
                        content = content + doc_content[j].text
                    normal_type.append(doc_content[i].text)
                    normal_content.append(content)
                    # 获取 药理毒理 开始节点
                    i_1 = i + 1
            if doc_content[i].text == '【执行标准】':
                normal_type.append(doc_content[i].text)
                normal_content.append(doc_content[i + 1].text)
            if doc_content[i].text == '【企业名称】' or doc_content[i].text == '【生产企业】':
                normal_type.append(doc_content[i].text)
                content_1 = doc_content[i + 1].text
                content_array_1 = content_1.split('：')
                if len(content_array_1) > 1:
                    normal_content.append(content_array_1[1])
                elif len(content_array_1) == 1:
                    content_array_1 = content_1.split(': ')
                    if len(content_array_1) > 1:
                        normal_content.append(content_array_1[1])
                    else:
                        normal_content.append(content_array_1)
                else:
                    normal_content.append(content_1)
def data_to_excel():
    data = iterate_files()
    # file_name = 'H20170167sms.doc'
    # data = decode_docx(file_name)
    # 新建excel表
    workbook = xlsxwriter.Workbook('/Users/heweiwen/Downloads/work/work/Python/国家药品监督管理局说明书_解析Excel_new.xls')
    # 新建sheet（sheet的名称为"sheet1"）
    worksheet = workbook.add_worksheet('国家药品监督管理局说明书')
    # 设置表头
    headings = ['wordName','通用名称','商品名称','英文名称','适应症','不良反应','注意事项','禁忌','药物过量','药理毒理','执行标准','企业名称']
    worksheet.write_row('A1', headings)
    i = 2
    for row in data:
        i = i + 1
        try:
            worksheet.write_row('A' + str(i), row)
        except KeyError:
            list_error.append(row)
            print('Error: worksheetwrite_row失败')
            continue

    # 将excel文件保存关闭，
    workbook.close()
    print('保存完毕')
    for path in not_files:
        print(path)
    for path in not_isdir:
        print(path)
    for path in list_error:
        print(path)

def download_excel():
    data = iterate_files()
    # file_name = 'gyzzH20203422sms.doc'
    # data = decode_docx(file_name)
    workbook = xlwt.Workbook('utf-8')
    sheet = workbook.add_sheet('国家药品监督管理局说明书')
    col = ['wordName','通用名称','商品名称','英文名称','适应症','不良反应','注意事项','禁忌','药物过量','药理毒理','执行标准','企业名称']
    for i in range(0, len(col)):
        sheet.write(0, i, col[i])
    for i in range(0, len(data)):
        new_data = data[i]
        for j in range(0, len(new_data)):
            try:
                sheet.write(i + 1, j,new_data[j])
            except KeyError:
                list_error.append(new_data[j])
                print('Error: worksheetwrite_row失败')
                continue

    workbook.save('/Users/heweiwen/Downloads/work/work/Python/国家药品监督管理局说明书_解析Excel_new.xls')
    print('保存完毕')
    for path in not_files:
        print(path)
    for path in not_isdir:
        print(path)
    for path in list_error:
        print(path)
if __name__ == '__main__':
    # iterate_files()
    # file_name = '2019S00318sms.doc'
    # data = decode_docx(file_name)
    download_excel()
    # data_to_excel()